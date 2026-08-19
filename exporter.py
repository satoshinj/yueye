# -*- coding: utf-8 -*-
"""输出格式转换: 将抓取结果导出为 PDF / Word / Markdown / 文本 / 图片合集。

PDF 走无损路径：抓取得到的原始 PNG 字节直接交给 img2pdf，不做中间的
PIL 解码 + JPEG 重编码（旧实现每页多一次有损压缩，白白损失清晰度）。
"""
from __future__ import annotations

import io
import re
from pathlib import Path

import img2pdf


def slugify(title: str) -> str:
    """生成安全的文件名。"""
    s = re.sub(r"[\\/:*?\"<>|]", "_", title)
    s = re.sub(r"\s+", "_", s).strip("_")
    return s[:80] or "document"


def _strip_alpha(data: bytes) -> bytes:
    """去掉 PNG 的 alpha 通道。

    canvas.toDataURL('image/png') 固定输出 RGBA，img2pdf 会为每页额外生成
    软掩码(SMask)，白白撑大 PDF。抓取时已把内容合成到白底，此处 RGBA->RGB
    是无损的。非 PNG 或无 alpha 则原样返回。
    """
    if not data[:8].startswith(b"\x89PNG"):
        return data
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(data))
        if img.mode not in ("RGBA", "LA", "P"):
            return data
        img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="PNG", optimize=True)
        return buf.getvalue()
    except Exception:
        return data


def _page_images(result) -> list[bytes]:
    """取出所有页面的图像字节（PDF 用，已去 alpha）。"""
    return [_strip_alpha(p.screenshot) for p in result.pages if p.screenshot]


# ---------------------------------------------------------------------------
def export_pdf(result, out_path: Path) -> Path:
    """导出 PDF。

    文章线路已由浏览器渲染出文字型 PDF（可选中、可搜索），直接落盘；
    阅读器/截图线路则把页面原图拼成图片型 PDF（无重编码）。
    """
    if getattr(result, "pdf_bytes", None):
        out_path.write_bytes(result.pdf_bytes)
        return out_path

    images = _page_images(result)
    if not images:
        raise ValueError("没有可导出的页面图像")
    out_path.write_bytes(img2pdf.convert(images))
    return out_path


# ---------------------------------------------------------------------------
def export_images(result, out_dir: Path) -> list[Path]:
    """导出为图片合集（原样落盘，不重编码）。"""
    if not any(p.screenshot for p in result.pages):
        # 文章线路产出的是文字型 PDF，没有页面图像。宁可明确报错，
        # 也不要静默生成 0 个文件让用户以为导出成功了。
        raise ValueError(
            "本次抓取走的是「网页文章」线路，没有页面图像，无法导出图片合集。\n"
            "请改用 PDF / Word / Markdown / 纯文本，"
            "或把线路手动指定为「整页截图」后重新抓取。")
    out_dir.mkdir(parents=True, exist_ok=True)
    paths = []
    for page in result.pages:
        if not page.screenshot:
            continue
        ext = ".png" if page.screenshot[:8].startswith(b"\x89PNG") else ".jpg"
        p = out_dir / f"page_{page.index + 1:04d}{ext}"
        p.write_bytes(page.screenshot)
        paths.append(p)
    return paths


# ---------------------------------------------------------------------------
def export_text(result, out_path: Path) -> Path:
    """导出纯文字。"""
    out_path.write_text(result.full_text, encoding="utf-8")
    return out_path


# ---------------------------------------------------------------------------
def export_markdown(result, out_path: Path) -> Path:
    """导出 Markdown（附带图片引用）。"""
    img_dir = out_path.with_suffix(".assets")
    rel_dir = img_dir.name
    paths = export_images(result, img_dir)

    md = [f"# {result.title}", "", f"> 来源: {result.url}", ""]
    if result.total_pages:
        md.append(f"> 抓取 {result.page_count} / {result.total_pages} 页")
    if result.stopped_reason:
        md.append(f"> 提前结束: {result.stopped_reason}")
    md += ["", result.full_text, "", "## 页面图像"]
    for i, p in enumerate(paths):
        md.append(f"![第{i + 1}页]({rel_dir}/{p.name})")
    out_path.write_text("\n".join(md), encoding="utf-8")
    return out_path


# ---------------------------------------------------------------------------
def export_word(result, out_path: Path) -> Path:
    """导出为 .docx（图片逐页插入；有文字则附文字）。

    旧实现用 RTF + \\ansi，中文必乱码；这里改用 python-docx。
    """
    from docx import Document
    from docx.shared import Cm

    doc = Document()
    doc.add_heading(result.title or "文档", level=1)
    doc.add_paragraph(f"来源: {result.url}")
    if result.total_pages:
        doc.add_paragraph(f"抓取 {result.page_count} / {result.total_pages} 页")
    if result.stopped_reason:
        doc.add_paragraph(f"提前结束: {result.stopped_reason}")

    if result.full_text.strip():
        doc.add_heading("文字内容", level=2)
        for line in result.full_text.split("\n"):
            if line.strip():
                doc.add_paragraph(line)

    images = _page_images(result)
    if images:
        doc.add_heading("页面图像", level=2)
        for img in images:
            doc.add_picture(io.BytesIO(img), width=Cm(16))

    doc.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
EXPORTERS = {
    "pdf": export_pdf,
    "images": export_images,
    "text": export_text,
    "markdown": export_markdown,
    "word": export_word,
}

EXTENSIONS = {"pdf": ".pdf", "text": ".txt", "markdown": ".md", "word": ".docx"}


def export(result, fmt: str, out_dir: Path) -> list[Path]:
    """按指定格式导出, 返回生成的文件路径列表。"""
    out_dir.mkdir(parents=True, exist_ok=True)
    base = slugify(result.title)

    fn = EXPORTERS.get(fmt)
    if fn is None:
        raise ValueError(f"不支持的格式: {fmt}")

    if fmt == "images":
        return fn(result, out_dir / base)
    return [fn(result, out_dir / (base + EXTENSIONS[fmt]))]
